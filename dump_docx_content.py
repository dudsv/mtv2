from docx import Document
import os

def dump_docx(docx_path):
    if not os.path.exists(docx_path):
        print(f"File not found: {docx_path}")
        return

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Error opening docx: {e}")
        return
    
    output_txt = "docx_dump.txt"
    with open(output_txt, "w", encoding="utf-8") as f:
        f.write(f"--- Content of {docx_path} ---\n\n")
        
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                f.write(f"P{i} [{p.style.name}]: {text}\n")
        
        for i, table in enumerate(doc.tables):
            f.write(f"\n--- Table {i} ---\n")
            for r_idx, row in enumerate(table.rows):
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                f.write(f"R{r_idx}: {row_text}\n")
    
    print(f"Dumped content to {output_txt}")

if __name__ == "__main__":
    base_folder = os.path.dirname(os.path.abspath(__file__))
    target_file = os.path.join(base_folder, "simulation_output", "simulation_result_v7.docx")
    dump_docx(target_file)
