"""
CrawlerThread worker for web crawling operations.
Fetches URLs, extracts specified data, and saves to Excel files.
"""

import os
import re
import datetime
import asyncio
import aiohttp
import openpyxl
from urllib.parse import urlparse
from bs4 import BeautifulSoup, Tag, NavigableString
from PyQt6.QtCore import QThread, pyqtSignal
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from config import MAX_EXCEL_CELL_LENGTH, TIMEOUT_STANDARD


class CrawlerThread(QThread):
    """
    A worker thread for crawling websites asynchronously.
    Fetches URLs, extracts specified data, and saves it to Excel files.
    """
    progress_update = pyqtSignal(int)
    log_update = pyqtSignal(str)
    finished = pyqtSignal(str)

    def __init__(self, mode, search_input, urls, extract_options, check_errors, output_folder):
        super().__init__()
        self.mode = mode  # 1: Search modules (classes), 2: Search words, 0: No search
        self.search_input = search_input
        self.urls = urls
        self.extract_options = extract_options
        self.check_errors = check_errors
        self.output_folder = output_folder or f"web_crawler_results_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        os.makedirs(self.output_folder, exist_ok=True)
        self.total_pages_crawled = 0
        self.stopped = False

    def stop(self):
        """Signals the thread to stop processing."""
        self.log_update.emit("Stopping crawler...")
        self.stopped = True

    def run(self):
        """Runs the asynchronous crawling process."""
        try:
            asyncio.run(self.main())
        except asyncio.CancelledError:
            self.log_update.emit("Crawling was cancelled")
        except Exception as e:
            self.log_update.emit(f"An unexpected error occurred: {e}")

    async def main(self):
        """Main async function to set up and execute crawling tasks."""
        class_patterns = []
        search_patterns = []

        if self.mode == 1 and self.search_input:
            class_patterns = [pattern.strip() for pattern in self.search_input.split(',') if pattern.strip()]
        elif self.mode == 2 and self.search_input:
            search_patterns = self._generate_search_patterns([word.strip() for word in self.search_input.split(',') if word.strip()])

        main_filename = os.path.join(self.output_folder, "results.xlsx")
        error_filename = os.path.join(self.output_folder, "error_results.xlsx") if self.check_errors else None
        main_filename = os.path.join(self.output_folder, "results.xlsx")
        error_filename = os.path.join(self.output_folder, "error_results.xlsx") if self.check_errors else None

        # "Extract All Meta" behaves like checking all SEO boxes
        if self.extract_options.get("meta_tags"):
            self.extract_options["title"] = True
            self.extract_options["h1"] = True
            self.extract_options["meta_description"] = True
            self.extract_options["og_tags"] = True
            self.extract_options["schema"] = True

        # Check if any Excel extraction options are selected
        excel_options_selected = any([
            self.extract_options.get("h1"),
            self.extract_options.get("title"),
            self.extract_options.get("meta_description"),
            self.extract_options.get("og_tags"),
            self.extract_options.get("schema"),
            self.mode in [1, 2]  # Search modes also require Excel
        ])

        # Only create Excel if needed
        wb_main = None
        ws_main = None
        
        if excel_options_selected:
            wb_main = openpyxl.Workbook()
            ws_main = wb_main.active
            ws_main.title = "Main Results"
            headers = ["URL"]
            
            # Reordered headers: H1 first, then Meta Title (renamed from Page Title)
            if self.extract_options.get("h1"): headers.append("H1 Tag")
            if self.extract_options.get("title"): headers.append("Meta Title") # Renamed from Page Title
            if self.extract_options.get("meta_description"): headers.append("Meta Description")
            if self.extract_options.get("og_tags"): 
                headers.append("OG Title")
                headers.append("OG Description")
                headers.append("OG Image")
            if self.extract_options.get("schema"): headers.append("Schema JSON")
            # Removed "All Meta Tags" blob column as requested
            
            if self.mode == 1: headers.append("Module Found")
            elif self.mode == 2: headers.append("Found Words")
            ws_main.append(headers)

        wb_errors = openpyxl.Workbook() if error_filename else None
        ws_errors = wb_errors.active if wb_errors else None
        if ws_errors:
            ws_errors.title = "Error Results"
            ws_errors.append(["URL", "Status Code", "Redirect"])

        async with aiohttp.ClientSession() as session:
            all_urls_to_check = []
            for url in self.urls:
                if self.stopped: break
                if urlparse(url).path.endswith(".xml"):
                    self.log_update.emit(f"Fetching URLs from sitemap: {url}")
                    sitemap_urls = await self._get_sitemap_urls(url, session)
                    all_urls_to_check.extend(sitemap_urls)
                else:
                    all_urls_to_check.append(url)
            
            total_urls = len(all_urls_to_check)
            self.total_pages_crawled = total_urls
            
            tasks = [self._crawl_url(u, session, class_patterns, search_patterns) for u in all_urls_to_check]
            
            for i, future in enumerate(asyncio.as_completed(tasks)):
                if  self.stopped: break
                
                result = await future
                if result:
                    if result['type'] == 'success':
                        if ws_main:  # Only append if Excel is being used
                            ws_main.append(result['main_data'])
                    elif result['type'] == 'error' and ws_errors:
                        ws_errors.append(result['error_data'])

                progress = int((i + 1) / total_urls * 100) if total_urls > 0 else 100
                self.progress_update.emit(progress)
                self.log_update.emit(f"Processed {i + 1}/{total_urls} URLs")

        if not self.stopped:
            if wb_main:  # Only save if Excel was created
                self.log_update.emit("Saving results to Excel files...")
                wb_main.save(main_filename)
            if wb_errors: wb_errors.save(error_filename)
            
            self.log_update.emit(f"Crawling completed. Results saved to {self.output_folder}")
            self.log_update.emit(f"Total pages processed: {self.total_pages_crawled}")
            self.finished.emit(self.output_folder)
        else:
            self.log_update.emit("Crawling stopped by user.")

    def _generate_search_patterns(self, words):
        return [re.compile(re.escape(word), re.IGNORECASE) for word in words]

    async def _get_sitemap_urls(self, sitemap_url, session):
        try:
            async with session.get(sitemap_url, ssl=False) as response:
                if response.status == 200:
                    content = await response.text()
                    soup = BeautifulSoup(content, 'lxml-xml')
                    return [loc.text for loc in soup.find_all('loc')]
                else:
                    self.log_update.emit(f"Sitemap fetch failed for {sitemap_url}: Status {response.status}")
                    return []
        except aiohttp.ClientError as e:
            self.log_update.emit(f"Network error reading sitemap {sitemap_url}: {e}")
            return []
        except Exception as e:
            self.log_update.emit(f"Error reading sitemap {sitemap_url}: {e}")
            return []

    async def _crawl_url(self, url, session, class_patterns, search_patterns):
        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            }
            async with session.get(url, ssl=False, timeout=TIMEOUT_STANDARD, headers=headers) as response:
                if response.status in {403, 404} and self.check_errors:
                    redirect_url = response.headers.get("Location", "N/A")
                    self.log_update.emit(f"Error {response.status} for {url}")
                    return {'type': 'error', 'error_data': [url, response.status, redirect_url]}

                if response.status == 200:
                    html = await response.text()
                    soup = BeautifulSoup(html, 'lxml')
                    result = {'type': 'success', 'url': url, 'soup': soup}
                    row_data = [url]
                    
                    # 1. H1 Tag (<h1>) - Now First
                    if self.extract_options.get("h1"):
                        h1 = soup.find("h1")
                        row_data.append(h1.get_text(strip=True) if h1 else "No H1")

                    # 2. Page Title (<title>) - Renamed to Meta Title
                    if self.extract_options.get("title"):
                        row_data.append(soup.title.string.strip() if soup.title and soup.title.string else "No title")
                    
                    # 3. Meta Description (Standard <meta name="description">)
                    if self.extract_options.get("meta_description"):
                        meta = soup.find("meta", attrs={"name": "description"})
                        row_data.append(meta["content"] if meta and meta.get("content") else "No meta description")
                    
                    # 4. OG Tags (Title, Description, Image)
                    if self.extract_options.get("og_tags"):
                        og_title = soup.find("meta", property="og:title")
                        og_desc = soup.find("meta", property="og:description")
                        og_image = soup.find("meta", property="og:image")
                        
                        og_t_val = og_title["content"] if og_title and og_title.get("content") else ""
                        og_d_val = og_desc["content"] if og_desc and og_desc.get("content") else ""
                        og_i_val = og_image["content"] if og_image and og_image.get("content") else ""
                        
                        row_data.append(og_t_val)
                        row_data.append(og_d_val)
                        row_data.append(og_i_val)
                    
                    # 5. Schema JSON-LD
                    if self.extract_options.get("schema"):
                        schemas = soup.find_all("script", type="application/ld+json")
                        if schemas:
                            # Join multiple schemas with a separator
                            schema_texts = [s.string.strip() for s in schemas if s.string]
                            full_schema = "\n---\n".join(schema_texts)
                            # Truncate if too long for Excel
                            row_data.append(full_schema[:MAX_EXCEL_CELL_LENGTH])
                        else:
                            row_data.append("No Schema")

                    # Removed "All Meta Tags" blob extraction

                    # Removed "All Meta Tags" blob extraction
                        
                    # 7. Search Modes
                    if self.mode == 1:
                        # Improved: Search in ANY tag, not just div
                        found = any(soup.find(class_=cp) for cp in class_patterns)
                        row_data.append("Yes" if found else "No")
                    elif self.mode == 2:
                        text = soup.get_text()
                        found_words = [p.pattern for p in search_patterns if p.search(text)]
                        row_data.append(', '.join(found_words) if found_words else "None")

                    result['main_data'] = row_data
                    return result
                else:
                    self.log_update.emit(f"Non-200 status for {url}: {response.status}")

        except asyncio.TimeoutError:
            self.log_update.emit(f"Timeout processing {url}")
        except aiohttp.ClientError as e:
            self.log_update.emit(f"Network error for {url}: {e}")
        except Exception as e:
            self.log_update.emit(f"Failed to process {url}: {e}")
        return None

    def _save_to_docx(self, url, soup, output_folder):
        """
        Generate a structured Word document from the crawled HTML.
        Detects Drupal blocks and preserves HTML hierarchy.
        """
        try:
            # Sanitize URL for filename
            safe_filename = re.sub(r'[<>:"/\\|?*]', '_', url)
            safe_filename = safe_filename[:100]  # Limit length
            docx_filename = os.path.join(output_folder, f"{safe_filename}.docx")
            
            doc = Document()
            
            # Add URL header
            title = doc.add_paragraph()
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title.add_run(f"Source: {url}")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 112, 192)
            
            doc.add_paragraph()  # Spacing
            
            # Process the main content area
            main_content = soup.find('main') or soup.find('body')
            
            if main_content:
                self._process_element(doc, main_content, level=0)
            else:
                doc.add_paragraph("No main content found.")
            
            doc.save(docx_filename)
            self.log_update.emit(f"Saved Word document: {safe_filename}.docx")
            
        except Exception as e:
            self.log_update.emit(f"Error saving Word document for {url}: {e}")
    
    def _process_element(self, doc, element, level=0):
        """
        Recursively process HTML elements and add to Word document.
        Detects Drupal blocks and maintains hierarchy.
        """
        # Skip script, style, nav, footer elements
        if element.name in ['script', 'style', 'nav', 'footer', 'header', 'noscript']:
            return
        
        # Detect Drupal blocks
        is_drupal_block = False
        block_id = None
        
        if element.name in ['div', 'section', 'article']:
            # Check for Drupal block markers
            classes = element.get('class', [])
            if 'block' in classes:
                is_drupal_block = True
                block_id = element.get('data-block-plugin-id') or element.get('id') or 'unknown'
            
            # Generic Component Detection
            # Matches: component--*, article--hero, hero--image
            component_name = None
            for cls in classes:
                if cls.startswith('component--'):
                    component_name = cls.replace('component--', '').replace('-', ' ').title()
                    break
                elif cls == 'article--hero':
                    component_name = "Hero Article"
                    break
                elif cls == 'hero--image':
                    component_name = "Hero Image"
                    break
            
            if component_name:
                self._process_special_component(doc, element, component_name)
                return  # Skip default processing for this component
        
        # If it's a Drupal block, add a marker
        if is_drupal_block:
            block_marker = doc.add_paragraph()
            run = block_marker.add_run(f"[DRUPAL BLOCK: {block_id}]")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.italic = True
        
        # Process based on tag type
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Headings
            heading_level = int(element.name[1])
            text = element.get_text(strip=True)
            if text:
                heading = doc.add_heading(text, level=heading_level)
        
        elif element.name == 'p':
            # Paragraphs
            text = element.get_text(strip=True)
            if text and len(text) > 1:  # Avoid empty paragraphs
                doc.add_paragraph(text)
        
        elif element.name in ['ul', 'ol']:
            # Lists
            for li in element.find_all('li', recursive=False):
                text = li.get_text(strip=True)
                if text:
                    doc.add_paragraph(text, style='List Bullet' if element.name == 'ul' else 'List Number')
        
        elif element.name == 'table':
            # Tables
            self._process_table(doc, element)
        
        else:
            # Process children for div, section, article, etc.
            for child in element.children:
                if isinstance(child, Tag):
                    self._process_element(doc, child, level + 1)
                elif isinstance(child, NavigableString):
                    text = str(child).strip()
                    if len(text) > 20:  # Only add significant text chunks
                        doc.add_paragraph(text)
    
    def _process_table(self, doc, table_element):
        """Process HTML table and add to Word document."""
        try:
            rows = table_element.find_all('tr')
            if not rows:
                return
            
            # Determine column count
            col_count = max(len(row.find_all(['td', 'th'])) for row in rows)
            
            # Create Word table
            word_table = doc.add_table(rows=len(rows), cols=col_count)
            word_table.style = 'Light Grid Accent 1'
            
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    if j < col_count:
                        text = cell.get_text(strip=True)
                        word_table.rows[i].cells[j].text = text
                        
                        # Bold header cells
                        if cell.name == 'th':
                            for paragraph in word_table.rows[i].cells[j].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        except Exception as e:
            self.log_update.emit(f"Error processing table: {e}")

    def _process_special_component(self, doc, element, component_name):
        """
        Generic handler for special components (Hero, Text-Image, Contact, etc.).
        Attempts to extract main image and text content and present them in a structured table.
        """
        try:
            # Add a marker
            marker = doc.add_paragraph()
            run = marker.add_run(f"[COMPONENT: {component_name}]")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.italic = True

            # Create 2-column table for structured layout
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # --- Extract Image ---
            img_tag = element.find('img')
            img_url = "No image found"
            if img_tag:
                img_url = img_tag.get('src', '')
                # Basic relative URL handling could go here
            
            # --- Extract Text ---
            # Heuristic: Get all text, but skip the image caption if it duplicates
            # We'll use a separator to join multiple text parts
            text_parts = []
            
            # Try specific text containers first
            text_containers = element.find_all(class_=re.compile(r'(summary|text|desc|body|content)'))
            if text_containers:
                for container in text_containers:
                    t = container.get_text(strip=True)
                    if t and len(t) > 10: text_parts.append(t)
            else:
                # Fallback: Get direct text from paragraphs or headings
                for tag in element.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    t = tag.get_text(strip=True)
                    if t: text_parts.append(t)
            
            # Deduplicate and join
            summary_text = "\n\n".join(list(dict.fromkeys(text_parts))) # simple dedup
            if not summary_text:
                summary_text = element.get_text(separator=' ', strip=True)[:500] # Ultimate fallback
            
            # --- Fill Table ---
            # Cell 0: Image Info
            cell_img = table.rows[0].cells[0]
            p_img = cell_img.paragraphs[0]
            if img_tag:
                p_img.add_run("IMAGE:\n").bold = True
                p_img.add_run(img_url)
            else:
                p_img.add_run("No Image").italic = True
            
            # Cell 1: Text Content
            cell_text = table.rows[0].cells[1]
            cell_text.text = summary_text
            
            doc.add_paragraph() # Spacing after component
            
        except Exception as e:
            self.log_update.emit(f"Error processing component {component_name}: {e}")
