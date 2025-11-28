"""
ContentExtractorWorker for structured Word document generation.
Fetches URL, extracts content with refined logic (V7.1), and saves to .docx.
"""

import os
import re
import asyncio
import aiohttp
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup, Tag, NavigableString
from PyQt6.QtCore import QThread, pyqtSignal
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from config import TIMEOUT_STANDARD


class ContentExtractorWorker(QThread):
    """
    Worker thread for extracting structured content from a URL and saving to Word.
    """
    log_update = pyqtSignal(str)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, url, output_folder):
        super().__init__()
        self.url = url
        self.output_folder = output_folder
        self.text_block_count = 0
        self.text_image_count = 0
        self.base_url = f"{urlparse(url).scheme}://{urlparse(url).netloc}"
        os.makedirs(self.output_folder, exist_ok=True)

    def run(self):
        """Run the async extraction process."""
        try:
            asyncio.run(self.extract())
        except Exception as e:
            self.error.emit(f"Error: {e}")
            self.log_update.emit(f"Extraction failed: {e}")

    async def extract(self):
        """Main async extraction function."""
        self.log_update.emit(f"Fetching {self.url}...")
        
        try:
            async with aiohttp.ClientSession() as session:
                headers = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
                }
                async with session.get(self.url, headers=headers, ssl=False, timeout=TIMEOUT_STANDARD) as response:
                    if response.status != 200:
                        self.error.emit(f"HTTP {response.status}")
                        self.log_update.emit(f"Failed to fetch URL: Status {response.status}")
                        return
                    
                    html = await response.text()
                    soup = BeautifulSoup(html, 'lxml')
                    
                    self.log_update.emit("Parsing content...")
                    docx_path = self._save_to_docx(self.url, soup, self.output_folder)
                    
                    self.log_update.emit(f"✓ Document saved: {os.path.basename(docx_path)}")
                    self.finished.emit(docx_path)
                    
        except asyncio.TimeoutError:
            self.error.emit("Request timeout")
            self.log_update.emit("Request timed out")
        except Exception as e:
            self.error.emit(str(e))
            self.log_update.emit(f"Error during extraction: {e}")

    def _save_to_docx(self, url, soup, output_folder):
        """Generate structured Word document from HTML."""
        try:
            # Create safe filename
            safe_filename = re.sub(r'[<>:"/\\|?*]', '_', url)
            safe_filename = safe_filename[:100] + ".docx"
            docx_filename = os.path.join(output_folder, safe_filename)
            
            doc = Document()
            
            # Header
            title = doc.add_paragraph()
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title.add_run(f"Source: {url}")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 112, 192)
            doc.add_paragraph()

            # V3: H1 Extraction (Top)
            h1 = soup.find('h1')
            if h1:
                h1_text = h1.get_text(strip=True)
                doc.add_heading(h1_text, level=1)
                doc.add_paragraph()

            # Reset counters
            self.text_block_count = 0
            self.text_image_count = 0

            # Process Main Content
            main_content = soup.find('main') or soup.find('body')
            if main_content:
                self._process_element(doc, main_content)
            else:
                doc.add_paragraph("No main content found.")
            
            # Add SEO Metadata Table at the end
            doc.add_paragraph()  # Spacer
            self._add_seo_metadata_table(doc, soup)
            
            doc.save(docx_filename)
            return docx_filename
            
        except Exception as e:
            self.log_update.emit(f"Error saving document: {e}")
            raise

    def _make_absolute(self, url):
        """Convert relative URLs to absolute."""
        if not url:
            return ""
        if url.startswith('http'):
            return url
        return urljoin(self.base_url, url)

    def _process_element(self, doc, element, level=0):
        """Recursively process HTML elements."""
        # Skip unwanted tags
        if element.name in ['script', 'style', 'nav', 'footer', 'header', 'noscript']:
            return

        # V2/V3: Exclusions
        classes = element.get('class', [])
        class_set = set(classes) if classes else set()
        
        exclusions = {
            'hero--article-wrapper--items',
            'article-author-bottom',
            'nppe-feedback-article-form',
            'related-topics--label',
            'related-topics--links',
            'article--progressbar',
            'component--newsletter'
        }
        
        if not class_set.isdisjoint(exclusions):
            return
        
        # Additional ID-based exclusion
        if element.get('id') == 'nppe-feedback-article-form':
            return

        # V1: Text Block Tagging
        target_classes = {
            'clearfix', 'text-formatted', 'field', 'field--name-field-c-text',
            'field--type-text-long', 'field--label-hidden', 'field__item'
        }
        if target_classes.issubset(class_set):
            self.text_block_count += 1
            marker = doc.add_paragraph()
            run = marker.add_run(f"[COMPONENT: Text Block {self.text_block_count}]")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 100, 0)
            run.bold = True

        # Special Components
        if element.name in ['div', 'section', 'article']:
            component_name = None
            for cls in classes:
                if cls.startswith('component--'):
                    component_name = cls.replace('component--', '').replace('-', ' ').title()
                    break
                elif cls == 'article--hero':
                    component_name = "Hero Article"
                    break
                elif cls == 'hero--image':
                    component_name = "Hero Image"
                    break
            
            if component_name:
                self._process_special_component(doc, element, component_name)
                return

        # Processing Tags
        if element.name == 'h1':
            # Skip H1 (already extracted at top)
            return

        if element.name in ['h2', 'h3', 'h4', 'h5', 'h6']:
            heading_level = int(element.name[1])
            heading = doc.add_heading(level=heading_level)
            self._process_children_with_links(heading, element)

        elif element.name == 'p':
            p = doc.add_paragraph()
            self._process_children_with_links(p, element)

        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                self._process_children_with_links(p, li)

        elif element.name == 'table':
            self._process_table(doc, element)

        else:
            # Recursion for containers
            for child in element.children:
                if isinstance(child, Tag):
                    self._process_element(doc, child, level + 1)
                elif isinstance(child, NavigableString):
                    text = str(child).strip()
                    if len(text) > 20:
                        doc.add_paragraph(text)

    def _process_children_with_links(self, paragraph, element):
        """Process children while preserving links with formatting."""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    paragraph.add_run(text)
            
            elif isinstance(child, Tag):
                if child.name == 'a':
                    link_text = child.get_text()
                    link_url = self._make_absolute(child.get('href', ''))
                    
                    # Link Text (Blue, Underlined)
                    run = paragraph.add_run(link_text)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True
                    
                    # URL (Blue, Underlined)
                    if link_url:
                        run_url = paragraph.add_run(f" ({link_url})")
                        run_url.font.color.rgb = RGBColor(0, 0, 255)
                        run_url.font.underline = True
                
                elif child.name in ['strong', 'b']:
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name in ['em', 'i']:
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                else:
                    paragraph.add_run(child.get_text())

    def _process_table(self, doc, table_element):
        """Process HTML table."""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        col_count = max(len(row.find_all(['td', 'th'])) for row in rows)
        word_table = doc.add_table(rows=len(rows), cols=col_count)
        word_table.style = 'Light Grid Accent 1'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < col_count:
                    word_table.rows[i].cells[j].text = cell.get_text(strip=True)

    def _process_special_component(self, doc, element, component_name):
        """Handle special components (Hero, Text Image, etc.)."""
        
        # --- 1. Hero Article ---
        if component_name == "Hero Article":
            # Extract Category
            cat_elem = element.find(class_='hero--article-category')
            if cat_elem:
                cat_text = cat_elem.get_text(strip=True)
                cat_link = cat_elem.find('a')
                cat_url = self._make_absolute(cat_link.get('href', '')) if cat_link else ""
                
                p_cat = doc.add_paragraph()
                p_cat.add_run(f"Article category: {cat_text}").bold = True
                if cat_url:
                    p_cat.add_run(f" - {cat_url}").italic = True

            # Create Table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Find image specifically from hero--image class
            hero_image_container = element.find(class_='hero--image')
            img_tag = None
            img_url = "No image found"
            img_alt = ""
            
            if hero_image_container:
                img_tag = hero_image_container.find('img')
                if img_tag:
                    img_url = self._make_absolute(img_tag.get('src', ''))
                    img_alt = img_tag.get('alt', '').strip()
            
            # Column 1: Image URL
            cell_img = table.rows[0].cells[0]
            p_img = cell_img.paragraphs[0]
            p_img.add_run("IMAGE:\n").bold = True
            p_img.add_run(img_url)
            
            # Column 2: Alt Text only
            cell_text = table.rows[0].cells[1]
            p_alt = cell_text.paragraphs[0]
            p_alt.add_run("Alt Text: ").bold = True
            p_alt.add_run(img_alt if img_alt else "No Alt Text")
            
            doc.add_paragraph()  # Spacing
            return

        # --- 2. Articles List ---
        if component_name == "Articles List":
            # Find component title
            component_title = element.find(['h2', 'h3', 'h4'])
            component_title_text = component_title.get_text(strip=True) if component_title else "Articles List"
            
            # Find all links in the component, excluding article-category
            all_links = element.find_all('a')
            # Filter out category links and deduplicate by URL
            seen_urls = set()
            article_links = []
            for link in all_links:
                parent = link.parent
                # Skip if parent has field--name-field-article-category class
                parent_classes = parent.get('class', []) if parent else []
                if 'field--name-field-article-category' not in parent_classes:
                    url = link.get('href', '')
                    # Deduplicate by URL (each article may have multiple links)
                    if url and url not in seen_urls:
                        seen_urls.add(url)
                        article_links.append(link)
            
            if article_links:
                # Create 2-column table: Title | URL
                # +1 row for header, +1 for component title row
                table = doc.add_table(rows=len(article_links) + 2, cols=2)
                table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = table.rows[0].cells
                p_title_header = header_cells[0].paragraphs[0]
                p_title_header.add_run("Title").bold = True
                p_url_header = header_cells[1].paragraphs[0]
                p_url_header.add_run("URL").bold = True
                
                # Component title row
                table.rows[1].cells[0].text = f"Component: {component_title_text}"
                table.rows[1].cells[1].text = ""
                # Make component title row bold
                for paragraph in table.rows[1].cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                
                # Data rows (articles)
                for i, link in enumerate(article_links, 2):  # Start at row 2 (after header and component title)
                    # Try to find the title field in parent elements
                    title_text = link.get_text(strip=True)  # Default to link text
                    
                    # Look for title field in parent hierarchy
                    parent = link.parent
                    while parent and parent != element:
                        title_field = parent.find(class_='field--name-title')
                        if title_field:
                            title_text = title_field.get_text(strip=True)
                            break
                        parent = parent.parent
                    
                    article_url = self._make_absolute(link.get('href', ''))
                    
                    # Use actual title instead of "Article N"
                    table.rows[i].cells[0].text = title_text
                    table.rows[i].cells[1].text = article_url
            
            doc.add_paragraph()  # Spacer
            return

        # --- 3. Text Image ---
        if component_name == "Text Image":
            self.text_image_count += 1
            display_name = f"{component_name} {self.text_image_count}"
            
            marker = doc.add_paragraph()
            run = marker.add_run(f"[COMPONENT: {display_name}]")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.italic = True

            # Create Table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Col 1: Image Info
            cell_info = table.rows[0].cells[0]
            p_info = cell_info.paragraphs[0]
            p_info.add_run("Image Position: Left\n").bold = True
            
            img_tag = element.find('img')
            alt_text = ""
            if img_tag:
                img_url = self._make_absolute(img_tag.get('src', ''))
                alt_text = img_tag.get('alt', '').strip()
                p_info.add_run("IMAGE:\n").bold = True
                p_info.add_run(img_url)
            else:
                p_info.add_run("No Image").italic = True
            
            # Col 2: Alt Text
            cell_alt = table.rows[0].cells[1]
            p_alt = cell_alt.paragraphs[0]
            p_alt.add_run("Alt Text: ").bold = True
            p_alt.add_run(alt_text if alt_text else "No Alt Text")
            
            # Content BELOW table
            doc.add_paragraph()  # Spacer
            
            # Tag-based extraction with formatting preservation
            content_tags = element.find_all(['p', 'h2', 'h3', 'h4', 'h5', 'h6', 'li'])
            
            # Track seen text to avoid duplicates
            seen_texts = set()
            
            if content_tags:
                for tag in content_tags:
                    if tag.find_parent(class_='hero--article-category'):
                        continue
                    
                    text_content = tag.get_text(strip=True)
                    if text_content and len(text_content) > 2 and text_content not in seen_texts:
                        seen_texts.add(text_content)
                        
                        # Add as heading or paragraph
                        if tag.name in ['h2', 'h3', 'h4', 'h5', 'h6']:
                            # Extract level from tag name (e.g., 'h3' -> 3)
                            try:
                                level = int(tag.name[1])
                            except ValueError:
                                level = 3 # Default to H3 if parsing fails
                            
                            doc.add_heading(text_content, level=level)
                        else:
                            doc.add_paragraph(text_content)
            else:
                # Fallback: simple text extraction
                text_containers = element.find_all(class_=re.compile(r'(summary|text|desc|body|content)'))
                if text_containers:
                    for container in text_containers:
                        t = container.get_text(strip=True)
                        if t and len(t) > 10 and t not in seen_texts:
                            seen_texts.add(t)
                            doc.add_paragraph(t)
            
            doc.add_paragraph()  # Spacing
            return

        # --- 4. Other Components (Default) ---
        # Standard 2-Column Table for other components
        marker = doc.add_paragraph()
        run = marker.add_run(f"[COMPONENT: {component_name}]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.italic = True

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        img_tag = element.find('img')
        img_url = "No image found"
        img_alt = ""
        if img_tag:
            img_url = self._make_absolute(img_tag.get('src', ''))
            img_alt = img_tag.get('alt', '').strip()
        
        # Standard text extraction
        text_parts = []
        text_containers = element.find_all(class_=re.compile(r'(summary|text|desc|body|content)'))
        if text_containers:
            for container in text_containers:
                if 'hero--article-category' in container.get('class', []):
                    continue
                t = container.get_text(strip=True)
                if t and len(t) > 10:
                    text_parts.append(t)
        
        if not text_parts:
            for tag in element.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                t = tag.get_text(strip=True)
                if t:
                    text_parts.append(t)
        
        summary_text = "\n\n".join(list(dict.fromkeys(text_parts)))
        if not summary_text:
            summary_text = element.get_text(separator=' ', strip=True)[:500]
        
        # Fill table
        cell_img = table.rows[0].cells[0]
        p_img = cell_img.paragraphs[0]
        if img_tag:
            p_img.add_run("IMAGE:\n").bold = True
            p_img.add_run(img_url)
        else:
            p_img.add_run("No Image").italic = True
        
        cell_text = table.rows[0].cells[1]
        cell_text.text = summary_text
        
        doc.add_paragraph()  # Spacing

    def _add_seo_metadata_table(self, doc, soup):
        """Add SEO metadata table at the end of the document."""
        try:
            # Add section header
            header = doc.add_paragraph()
            run = header.add_run("SEO METADATA")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 112, 192)
            
            # Extract URL path
            from urllib.parse import urlparse
            url_path = urlparse(self.url).path
            
            # Extract metadata
            meta_title = soup.find('title')
            meta_title_text = meta_title.string.strip() if meta_title and meta_title.string else "No Meta Title"
            
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            meta_desc_text = meta_desc.get('content', 'No Meta Description') if meta_desc else "No Meta Description"
            
            og_title = soup.find('meta', property='og:title')
            og_title_text = og_title.get('content', 'No OG Title') if og_title else "No OG Title"
            
            og_desc = soup.find('meta', property='og:description')
            og_desc_text = og_desc.get('content', 'No OG Description') if og_desc else "No OG Description"
            
            schemas = soup.find_all('script', type='application/ld+json')
            if schemas:
                schema_texts = [s.string.strip() for s in schemas if s.string]
                schema_text = "\n---\n".join(schema_texts)
                # Truncate if too long
                if len(schema_text) > 2000:
                    schema_text = schema_text[:2000] + "...[truncated]"
            else:
                schema_text = "No Schema"
            
            # Create 2-column table with 6 rows (one per metadata field)
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Define metadata rows
            metadata_rows = [
                ('Internal URL', url_path),
                ('Meta Title', meta_title_text),
                ('Meta Description', meta_desc_text),
                ('OG Title', og_title_text),
                ('OG Description', og_desc_text),
                ('Schema', schema_text)
            ]
            
            # Fill table
            for i, (label, value) in enumerate(metadata_rows):
                # Column 1: Label (bold)
                cell_label = table.rows[i].cells[0]
                p_label = cell_label.paragraphs[0]
                run_label = p_label.add_run(label)
                run_label.bold = True
                
                # Column 2: Value
                cell_value = table.rows[i].cells[1]
                cell_value.text = value
                
        except Exception as e:
            self.log_update.emit(f"Error adding SEO metadata table: {e}")
