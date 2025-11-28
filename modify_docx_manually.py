from docx import Document
from docx.shared import Pt, RGBColor
import os

FILE_PATH = r"c:\Users\duduv\OneDrive\Documentos\Coding\web_crawler_app\https___www.purina.fr_choisir-animal_articles_accueillir-chien_prenom_japonais.docx"

def add_marker(doc, paragraph, label):
    # Insert a new paragraph before the given paragraph
    # python-docx doesn't have a simple insert_paragraph_before for all cases, 
    # but we can try to insert it into the element tree or just add it and move it.
    # Actually, paragraph.insert_paragraph_before() exists!
    
    new_p = paragraph.insert_paragraph_before()
    run = new_p.add_run(f"[COMPONENT: {label}]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 100, 0) # Dark Green for visibility
    run.bold = True
    print(f"Added {label}")

def main():
    if not os.path.exists(FILE_PATH):
        print(f"File not found: {FILE_PATH}")
        return

    doc = Document(FILE_PATH)
    
    # Target text snippets (simplified for matching)
    targets = [
        ("Accueillir un chien chez soi est une expérience", "Text Block 1"),
        ("10 noms de chiens japonais les plus populaires", "Text Block 2"), # Matches Heading 2
        ("Kotaro (petit garçon)", "Text Block 3"), # Matches List Item
        ("Dango (boulettes)", "Text Block 4") # Matches List Item
    ]
    
    found_counts = {t[1]: 0 for t in targets}
    
    # Iterate through paragraphs
    # We need to be careful modifying the list while iterating, but insert_paragraph_before is safe-ish
    # if we iterate over a copy or handle indices. 
    # Safest is to collect target paragraphs first.
    
    paragraphs_to_tag = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        for snippet, label in targets:
            if snippet in text and found_counts[label] == 0:
                paragraphs_to_tag.append((p, label))
                found_counts[label] += 1
                break # Only tag once per label
    
    for p, label in paragraphs_to_tag:
        add_marker(doc, p, label)
        
    new_file_path = FILE_PATH.replace(".docx", "_modified.docx")
    try:
        doc.save(new_file_path)
        print(f"Modified file saved: {new_file_path}")
    except PermissionError:
        print(f"Error: Could not save to {new_file_path}. Is it open?")


if __name__ == "__main__":
    main()
