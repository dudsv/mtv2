from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

FILE_PATH = r"c:\Users\duduv\OneDrive\Documentos\Coding\web_crawler_app\https___www.purina.fr_choisir-animal_articles_accueillir-chien_prenom_japonais_modified.docx"

def format_as_link(paragraph, text_to_find):
    if text_to_find in paragraph.text:
        # This is a simplistic replacement that might mess up runs if not careful,
        # but for a simulation/verification script it's often enough to just 
        # append a new run or try to find the run containing the text.
        # Since python-docx doesn't support easy "find and format substring",
        # we will do a naive replace: clear paragraph and re-add with formatting.
        # WARNING: This removes existing formatting (like bolding of other parts).
        # Given this is a specific test file, we can try to be slightly smarter 
        # or just accept the limitation for the validation.
        
        # Better approach for simulation: 
        # Just find the run? No, text might be split across runs.
        # Let's try to just highlight the WHOLE paragraph if it matches exactly, 
        # or split the text.
        
        # For the TOC items (which are list bullets), the whole text is the link.
        if paragraph.text.strip() == text_to_find.strip():
            # FIX: Do not format Headings as links
            if paragraph.style.name.startswith("Heading"):
                print(f"Skipping Heading: '{text_to_find}'")
                return False

            paragraph.clear()
            run = paragraph.add_run(text_to_find)
            run.font.color.rgb = RGBColor(0, 0, 255) # Blue
            run.font.underline = True
            print(f"Formatted full paragraph: '{text_to_find}'")
            return True
            
        # For embedded links (like "meilleurs prénoms..."), we need to split.
        if text_to_find in paragraph.text:
            original_text = paragraph.text
            parts = original_text.split(text_to_find)
            paragraph.clear()
            
            # Rebuild
            for i, part in enumerate(parts):
                paragraph.add_run(part)
                if i < len(parts) - 1:
                    run = paragraph.add_run(text_to_find)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True
            print(f"Formatted embedded link: '{text_to_find}'")
            return True
    return False

def main():
    try:
        doc = Document(FILE_PATH)
    except Exception as e:
        print(f"Error opening file: {e}")
        return

    targets = [
        "meilleurs prénoms de petits chiens",
        "10 noms de chiens japonais les plus populaires",
        "Top 10 des noms masculins de chiens au Japon",
        "Top 10 des noms féminins de chiens au Japon",
        "20 noms de chiens japonais craquants",
        "20 Noms de chiens japonais inspirés de la Pop culture",
        "20 noms de plats japonais pour chiens",
        "Conseils pour choisir un nom japonais idéal pour votre chien"
    ]
    
    count = 0
    for p in doc.paragraphs:
        for target in targets:
            if format_as_link(p, target):
                count += 1
                
    try:
        doc.save(FILE_PATH)
        print(f"Successfully updated {FILE_PATH}")
        print(f"Total links simulated: {count}")
    except PermissionError:
        print(f"Error: Could not save to {FILE_PATH}. Is it open?")

if __name__ == "__main__":
    main()
