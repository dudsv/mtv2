import sys
from docx import Document

def analyze_docx(file_path):
    print(f"Analyzing: {file_path}")
    try:
        doc = Document(file_path)
        
        print(f"\nTotal Paragraphs: {len(doc.paragraphs)}")
        print(f"Total Tables: {len(doc.tables)}")
        
        print("\n--- Content Overview ---")
        
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                style = p.style.name
                # Print first few chars of content
                content = p.text[:100] + "..." if len(p.text) > 100 else p.text
                print(f"P{i} [{style}]: {content}")
                
        print("\n--- Table Overview ---")
        for i, table in enumerate(doc.tables):
            print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
            # Print first row content
            if table.rows:
                first_row = [cell.text.strip() for cell in table.rows[0].cells]
                print(f"  Row 0: {first_row}")

    except Exception as e:
        print(f"Error reading file: {e}")

if __name__ == "__main__":
    # Use the specific file mentioned by the user
    file_path = r"c:\Users\duduv\OneDrive\Documentos\Coding\web_crawler_app\https___www.purina.fr_choisir-animal_articles_accueillir-chien_prenom_japonais.docx"
    
    # Redirect stdout to a file
    with open("analysis_result.txt", "w", encoding="utf-8") as f:
        sys.stdout = f
        analyze_docx(file_path)

