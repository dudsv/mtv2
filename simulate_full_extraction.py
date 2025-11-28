import asyncio
import aiohttp
import re
import os
from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configuration
OUTPUT_FOLDER = "simulation_output"
TARGET_URL = "https://www.purina.fr/choisir-animal/articles/accueillir-chien/prenom/japonais"
TARGET_CLASS_STR = "clearfix text-formatted field field--name-field-c-text field--type-text-long field--label-hidden field__item"
TARGET_CLASSES = set(TARGET_CLASS_STR.split())

class SimulationCrawler:
    def __init__(self):
        self.text_block_count = 0
        self.text_image_count = 0
        if not os.path.exists(OUTPUT_FOLDER):
            os.makedirs(OUTPUT_FOLDER)

    async def run(self):
        print(f"Fetching {TARGET_URL}...")
        async with aiohttp.ClientSession() as session:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            }
            async with session.get(TARGET_URL, headers=headers) as response:
                if response.status != 200:
                    print(f"Failed: {response.status}")
                    return
                html = await response.text()
                soup = BeautifulSoup(html, 'lxml')
                self._save_to_docx(TARGET_URL, soup, OUTPUT_FOLDER)

    def _save_to_docx(self, url, soup, output_folder):
        try:
            safe_filename = "simulation_result_v7.docx"
            docx_filename = os.path.join(output_folder, safe_filename)
            
            doc = Document()
            
            # Header
            title = doc.add_paragraph()
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title.add_run(f"Source: {url}")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 112, 192)
            doc.add_paragraph()

            # --- V3: H1 Extraction (Top) ---
            h1 = soup.find('h1')
            if h1:
                h1_text = h1.get_text(strip=True)
                doc.add_heading(h1_text, level=1)
                doc.add_paragraph() # Spacer

            # Reset counters
            self.text_block_count = 0
            self.text_image_count = 0

            # Process Main
            main_content = soup.find('main') or soup.find('body')
            if main_content:
                self._process_element(doc, main_content)
            else:
                doc.add_paragraph("No main content found.")
            
            doc.save(docx_filename)
            print(f"Saved: {docx_filename}")
            
        except Exception as e:
            print(f"Error saving docx: {e}")
            import traceback
            traceback.print_exc()

    def _process_element(self, doc, element, level=0):
        if element.name in ['script', 'style', 'nav', 'footer', 'header', 'noscript']:
            return

        # --- V2/V3: Exclusions ---
        classes = element.get('class', [])
        class_set = set(classes) if classes else set()
        
        exclusions = {
            'hero--article-wrapper--items',
            'article-author-bottom',
            'nppe-feedback-article-form',
            'related-topics--label',
            'related-topics--links',
            'article--progressbar'
        }
        # Check if any exclusion class is present
        if not class_set.isdisjoint(exclusions):
            return
        
        # Specific component exclusion
        if 'component--newsletter' in class_set:
            return

        # V3: Stronger Feedback Exclusion (Check ID or Children)
        if element.get('id') == 'nppe-feedback-article-form':
            return
        
        # Check if this element IS the feedback form container by text content if class missing
        if element.name == 'div' and "Aidez-nous à nous améliorer" in element.get_text() and len(element.get_text()) < 1000:
             # Double check if it's the specific feedback block
             if element.find(string="Aidez-nous à nous améliorer"):
                 # But be careful not to skip the whole body if it's a parent.
                 # Only skip if it's relatively small or specific structure.
                 # We'll assume the class check above works for the main container, 
                 # this is a fallback for inner containers if they lack classes.
                 # Actually, let's be conservative: if it has the text and is a small block, skip.
                 if len(element.get_text()) < 300:
                     return

        # --- Logic 1: Text Block Tagging ---
        # Check for specific Text Block class
        if TARGET_CLASSES.issubset(class_set):
            self.text_block_count += 1
            marker = doc.add_paragraph()
            run = marker.add_run(f"[COMPONENT: Text Block {self.text_block_count}]")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 100, 0) # Dark Green
            run.bold = True
            # Continue processing children!

        # Check for Drupal Block (V2: Markers Removed)
        # We still need to detect special components inside blocks
        if element.name in ['div', 'section', 'article']:
            # Special Components
            component_name = None
            for cls in classes:
                if cls.startswith('component--'):
                    component_name = cls.replace('component--', '').replace('-', ' ').title()
                    break
                elif cls == 'article--hero':
                    component_name = "Hero Article"
                    break
                elif cls == 'hero--image':
                    component_name = "Hero Image"
                    break
            
            if component_name:
                self._process_special_component(doc, element, component_name)
                return

        # --- Processing Tags ---
        if element.name == 'h1':
            # V3: Skip H1 here as it's extracted at the top
            return

        if element.name in ['h2', 'h3', 'h4', 'h5', 'h6']:
            heading_level = int(element.name[1])
            heading = doc.add_heading(level=heading_level)
            self._process_children_with_links(heading, element)

        elif element.name == 'p':
            p = doc.add_paragraph()
            self._process_children_with_links(p, element)

        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                self._process_children_with_links(p, li)

        elif element.name == 'table':
            self._process_table(doc, element)

        else:
            # Recursion for containers
            for child in element.children:
                if isinstance(child, Tag):
                    self._process_element(doc, child, level + 1)
                elif isinstance(child, NavigableString):
                    text = str(child).strip()
                    if len(text) > 20:
                        doc.add_paragraph(text)

    def _make_absolute(self, url):
        if url.startswith('http'):
            return url
        if url.startswith('/'):
            # Simple join for this specific site
            base = "https://www.purina.fr"
            return base + url
        return url

    def _process_children_with_links(self, paragraph, element):
        """
        Iterates over children to preserve <a> tags with formatting and URL.
        """
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    paragraph.add_run(text)
            
            elif isinstance(child, Tag):
                if child.name == 'a':
                    link_text = child.get_text()
                    link_url = self._make_absolute(child.get('href', ''))
                    
                    # Add Link Text (Blue, Underlined)
                    run = paragraph.add_run(link_text)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True
                    
                    # Add URL (Blue, Underlined)
                    if link_url:
                        run_url = paragraph.add_run(f" ({link_url})")
                        run_url.font.color.rgb = RGBColor(0, 0, 255)
                        run_url.font.underline = True
                
                elif child.name in ['strong', 'b']:
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name in ['em', 'i']:
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                else:
                    paragraph.add_run(child.get_text())

    def _process_table(self, doc, table_element):
        # (Simplified copy of existing logic)
        rows = table_element.find_all('tr')
        if not rows: return
        col_count = max(len(row.find_all(['td', 'th'])) for row in rows)
        word_table = doc.add_table(rows=len(rows), cols=col_count)
        word_table.style = 'Light Grid Accent 1'
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < col_count:
                    word_table.rows[i].cells[j].text = cell.get_text(strip=True)

    def _process_special_component(self, doc, element, component_name):
        # V3: Extract Category from Hero
        if component_name == "Hero Article":
            cat_elem = element.find(class_='hero--article-category')
            if cat_elem:
                cat_text = cat_elem.get_text(strip=True)
                cat_link = cat_elem.find('a')
                cat_url = self._make_absolute(cat_link.get('href', '')) if cat_link else ""
                
                p_cat = doc.add_paragraph()
                p_cat.add_run(f"Article category: {cat_text}").bold = True
                if cat_url:
                    p_cat.add_run(f" - {cat_url}").italic = True

        # V4: Numbering for Text Image
        display_name = component_name
        if component_name == "Text Image":
            self.text_image_count += 1
            display_name = f"{component_name} {self.text_image_count}"

        marker = doc.add_paragraph()
        run = marker.add_run(f"[COMPONENT: {display_name}]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.italic = True

        # V7: Text Image uses 2 Columns + Text Below
        if component_name == "Text Image":
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Col 1: Image Info
            cell_info = table.rows[0].cells[0]
            p_info = cell_info.paragraphs[0]
            p_info.add_run("Image Position: Left\n").bold = True
            
            img_tag = element.find('img')
            img_url = "No image found"
            alt_text = ""
            if img_tag:
                img_url = self._make_absolute(img_tag.get('src', ''))
                alt_text = img_tag.get('alt', '').strip()
                p_info.add_run("IMAGE:\n").bold = True
                p_info.add_run(img_url)
            else:
                p_info.add_run("No Image").italic = True
            
            # Col 2: Alt Text
            cell_alt = table.rows[0].cells[1]
            p_alt = cell_alt.paragraphs[0]
            p_alt.add_run("Alt Text: ").bold = True
            p_alt.add_run(alt_text if alt_text else "No Alt Text")
            
            # Content BELOW table
            doc.add_paragraph() # Spacer
            
            # Extraction Logic (Tag-based to avoid duplication)
            text_parts = []
            # Prioritize atomic content tags to avoid Parent-Child duplication
            content_tags = element.find_all(['p', 'h2', 'h3', 'h4', 'h5', 'h6', 'li'])
            
            if content_tags:
                for tag in content_tags:
                    # Skip if it's the "Previous/Next" block (often in a specific wrapper, but we clean text later)
                    # Also skip if it's inside the image container (unlikely for p tags but good to be safe)
                    if tag.find_parent(class_='hero--article-category'): continue
                    
                    t = tag.get_text(strip=True)
                    if t and len(t) > 2: # Filter noise
                        text_parts.append(t)
            else:
                # Fallback to class-based if no tags found (rare)
                text_containers = element.find_all(class_=re.compile(r'(summary|text|desc|body|content)'))
                if text_containers:
                    for container in text_containers:
                        t = container.get_text(strip=True)
                        if t and len(t) > 10: text_parts.append(t)
            
            # Deduplicate while preserving order
            unique_parts = []
            seen = set()
            for part in text_parts:
                if part not in seen:
                    unique_parts.append(part)
                    seen.add(part)

            summary_text = "\n\n".join(unique_parts)
            if not summary_text: summary_text = element.get_text(separator=' ', strip=True)[:500]
            
            # V6/V7: Pagination Cleanup
            for junk in ["Previous", "Next", "1sur1"]:
                summary_text = summary_text.replace(junk, "")
            
            # Add as paragraph below
            p_content = doc.add_paragraph(summary_text.strip())

        else:
            # Standard 2-Column Table for other components
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            img_tag = element.find('img')
            img_url = "No image found"
            if img_tag:
                img_url = self._make_absolute(img_tag.get('src', ''))
            
            text_parts = []
            if component_name == "Hero Article":
                # V3: Separate fields with newlines
                hero_fields = element.find_all(class_=re.compile(r'(hero--title|hero--subtitle|field--name-title|field--name-uid|field--name-created)'))
                for t in hero_fields:
                    text_parts.append(t.get_text(strip=True))
            
            # Fallback / Standard Text Extraction
            text_containers = element.find_all(class_=re.compile(r'(summary|text|desc|body|content)'))
            if text_containers:
                for container in text_containers:
                    if 'hero--article-category' in container.get('class', []): continue
                    t = container.get_text(strip=True)
                    if t and len(t) > 10: text_parts.append(t)
            
            if not text_parts:
                 for tag in element.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    t = tag.get_text(strip=True)
                    if t: text_parts.append(t)
            
            summary_text = "\n\n".join(list(dict.fromkeys(text_parts)))
            if not summary_text: summary_text = element.get_text(separator=' ', strip=True)[:500]
            
            cell_img = table.rows[0].cells[0]
            p_img = cell_img.paragraphs[0]
            if img_tag:
                p_img.add_run("IMAGE:\n").bold = True
                p_img.add_run(img_url)
            else:
                p_img.add_run("No Image").italic = True
            
            cell_text = table.rows[0].cells[1]
            cell_text.text = summary_text
        
        doc.add_paragraph()

if __name__ == "__main__":
    crawler = SimulationCrawler()
    asyncio.run(crawler.run())
