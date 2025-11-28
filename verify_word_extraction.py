import os
import shutil
from bs4 import BeautifulSoup
from workers.crawler_worker import CrawlerThread
from docx import Document

# Mock CrawlerThread to avoid full init
class MockCrawler(CrawlerThread):
    def __init__(self):
        # Skip super init to avoid QThread/Signal issues in simple script
        self.log_update = self
        self.output_folder = "test_output"
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)

    def emit(self, msg):
        print(f"[LOG] {msg}")

def test_word_extraction():
    print("Testing Word Extraction Logic...")
    
    # 1. Create Mock HTML
    html_content = """
    <html>
    <body>
        <main>
            <h1>Page Title</h1>
            <p>Introductory paragraph.</p>
            
            <!-- Drupal Block -->
            <div class="block" data-block-plugin-id="block_123">
                <h2>Block Title</h2>
                <p>Block content.</p>
            </div>
            
            <!-- Special Component: Text Image -->
            <div class="component component--text-image">
                <img src="https://example.com/image.jpg" alt="Test Image">
                <div class="field--name-body">
                    <p>Component description text.</p>
                </div>
            </div>
            
            <!-- Standard Table -->
            <table>
                <tr><th>Header 1</th><th>Header 2</th></tr>
                <tr><td>Row 1 Col 1</td><td>Row 1 Col 2</td></tr>
            </table>
        </main>
    </body>
    </html>
    """
    
    soup = BeautifulSoup(html_content, 'lxml')
    
    # 2. Initialize Worker
    worker = MockCrawler()
    
    # 3. Run Extraction
    output_file = os.path.join(worker.output_folder, "test_extraction.docx")
    # Clean up previous run
    if os.path.exists(output_file):
        os.remove(output_file)
        
    print(f"Generating {output_file}...")
    worker._save_to_docx("https://example.com/test-page", soup, worker.output_folder)
    
    # 4. Verify Output
    if os.path.exists(output_file):
        print("SUCCESS: File created.")
        
        doc = Document(output_file)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        print("\n--- Document Content Check ---")
        checks = {
            "Source URL": "Source: https://example.com/test-page",
            "H1": "Page Title",
            "Drupal Block Marker": "[DRUPAL BLOCK: block_123]",
            "Component Marker": "[COMPONENT: Text Image]",
            "Table Content": "Row 1 Col 1" # Tables are in tables, not paragraphs, but let's check basic text if possible or just existence
        }
        
        for label, expected in checks.items():
            if expected in full_text:
                print(f"[PASS] Found {label}")
            else:
                # Tables are special, check tables list
                if label == "Table Content":
                    found_table = False
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if expected in cell.text:
                                    found_table = True
                    if found_table:
                        print(f"[PASS] Found {label} in table")
                    else:
                        print(f"[FAIL] Missing {label}")
                else:
                    print(f"[FAIL] Missing {label}: '{expected}'")

    else:
        print("FAIL: File not created.")

if __name__ == "__main__":
    test_word_extraction()
