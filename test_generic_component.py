"""
Test script for generic component export.
"""

import asyncio
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
import re

# Mocking the worker class methods
class MockWorker:
    def __init__(self):
        self.log_update = self
    
    def emit(self, msg):
        print(f"[LOG] {msg}")

    def _process_special_component(self, doc, element, component_name):
        # Copied logic for testing
        try:
            marker = doc.add_paragraph()
            run = marker.add_run(f"[COMPONENT: {component_name}]")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.italic = True

            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            img_tag = element.find('img')
            img_url = "No image found"
            if img_tag:
                img_url = img_tag.get('src', '')
            
            text_parts = []
            text_containers = element.find_all(class_=re.compile(r'(summary|text|desc|body|content)'))
            if text_containers:
                for container in text_containers:
                    t = container.get_text(strip=True)
                    if t and len(t) > 10: text_parts.append(t)
            else:
                for tag in element.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    t = tag.get_text(strip=True)
                    if t: text_parts.append(t)
            
            summary_text = "\n\n".join(list(dict.fromkeys(text_parts)))
            if not summary_text:
                summary_text = element.get_text(separator=' ', strip=True)[:500]
            
            cell_img = table.rows[0].cells[0]
            p_img = cell_img.paragraphs[0]
            if img_tag:
                p_img.add_run("IMAGE:\n").bold = True
                p_img.add_run(img_url)
            else:
                p_img.add_run("No Image").italic = True
            
            cell_text = table.rows[0].cells[1]
            cell_text.text = summary_text
            
            doc.add_paragraph()
            print(f"Processed {component_name}")
            
        except Exception as e:
            print(f"Error: {e}")

async def test_generic():
    html = """
    <div>
        <!-- 1. Text-Image Component -->
        <div class="component component--text-image">
            <img src="dog.jpg">
            <div class="field--name-field-c-sideimagetext-summary">Text Image Summary</div>
        </div>

        <!-- 2. Hero Article -->
        <div class="article--hero">
            <img src="hero.jpg">
            <h1>Hero Title</h1>
            <p>Hero description text.</p>
        </div>

        <!-- 3. Contact Us -->
        <div class="component component--contact-us-small">
            <h3>Contact Us</h3>
            <p>Call us at 123-456-7890</p>
        </div>
    </div>
    """
    
    soup = BeautifulSoup(html, 'lxml')
    doc = Document()
    worker = MockWorker()
    
    # Simulate detection loop
    for element in soup.find_all(class_=True):
        classes = element.get('class', [])
        component_name = None
        
        for cls in classes:
            if cls.startswith('component--'):
                component_name = cls.replace('component--', '').replace('-', ' ').title()
                break
            elif cls == 'article--hero':
                component_name = "Hero Article"
                break
            elif cls == 'hero--image':
                component_name = "Hero Image"
                break
        
        if component_name:
            worker._process_special_component(doc, element, component_name)

    doc.save("test_generic.docx")
    print("Saved test_generic.docx")

if __name__ == "__main__":
    asyncio.run(test_generic())
